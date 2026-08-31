#!/usr/bin/env python3
"""Build the Inzint proposal for Taif Al-Kaabi.

The document uses the Inzint Oman visual system and deliberately separates
implemented SaaS capabilities from readiness-gated or separately scoped work.
"""

from __future__ import annotations

import os
from datetime import datetime, timezone
from pathlib import Path

from PIL import Image, ImageDraw
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "output" / "proposal"
TMP_DIR = ROOT / "tmp" / "proposal_work"
DOCX_PATH = OUT_DIR / "Inzint_Proposal_Taif_Al-Kaabi_WhatsApp_AI_SaaS.docx"
LOGO_PATH = ROOT.parent / "inzint" / "oman-website" / "public" / "inzint-oman-light-text.png"
BANNER_PATH = TMP_DIR / "inzint_proposal_banner.png"

NAVY = "000C60"
NAVY_2 = "050D61"
PURPLE = "6200EE"
PURPLE_DARK = "4B00B5"
PURPLE_SOFT = "7A63DF"
PINK = "FF0266"
SURFACE = "F1F6FC"
SURFACE_2 = "F7F9FE"
WHITE = "FFFFFF"
INK = "151319"
BODY = "28314A"
MUTED = "667085"
BORDER = "D8D9E8"
GREEN = "18794E"
GREEN_BG = "EAF7F1"
AMBER = "8A5A00"
AMBER_BG = "FFF5DB"
LILAC_BG = "F1EBFF"
RED = "A13A3A"
RED_BG = "FCEEEE"

FONT = "Onest"
ARABIC_FONT = "Arial Unicode MS"
PAGE_WIDTH_DXA = 12240
PAGE_HEIGHT_DXA = 15840
MARGIN_DXA = 1440
CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_run_font(run, *, name=FONT, size=None, color=None, bold=None, italic=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), name)
    if size is not None:
        run.font.size = Pt(size)
        sz_cs = rpr.find(qn("w:szCs"))
        if sz_cs is None:
            sz_cs = OxmlElement("w:szCs")
            rpr.append(sz_cs)
        sz_cs.set(qn("w:val"), str(round(size * 2)))
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    return run


def set_para_spacing(p, *, before=0, after=0, line=1.0, keep=False, keep_next=False):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    pf.keep_together = keep
    pf.keep_with_next = keep_next


def set_paragraph_shading(p, fill: str):
    ppr = p._p.get_or_add_pPr()
    shd = ppr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        ppr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def set_paragraph_border(p, *, side="left", color=PURPLE, size=18, space=8):
    ppr = p._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    edge = OxmlElement(f"w:{side}")
    edge.set(qn("w:val"), "single")
    edge.set(qn("w:sz"), str(size))
    edge.set(qn("w:space"), str(space))
    edge.set(qn("w:color"), color)
    pbdr.append(edge)


def set_cell_shading(cell, fill: str):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, *, top=100, start=120, bottom=100, end=120):
    tcpr = cell._tc.get_or_add_tcPr()
    mar = tcpr.find(qn("w:tcMar"))
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tcpr.append(mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_borders(cell, *, color=BORDER, size=6, sides=("top", "left", "bottom", "right", "insideH", "insideV")):
    tcpr = cell._tc.get_or_add_tcPr()
    borders = tcpr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcpr.append(borders)
    for side in sides:
        edge = borders.find(qn(f"w:{side}"))
        if edge is None:
            edge = OxmlElement(f"w:{side}")
            borders.append(edge)
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), str(size))
        edge.set(qn("w:space"), "0")
        edge.set(qn("w:color"), color)


def set_cell_no_borders(cell):
    tcpr = cell._tc.get_or_add_tcPr()
    borders = tcpr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcpr.append(borders)
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge = OxmlElement(f"w:{side}")
        edge.set(qn("w:val"), "nil")
        borders.append(edge)


def set_table_geometry(table, widths_dxa, *, indent_dxa=TABLE_INDENT_DXA):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblpr = table._tbl.tblPr

    tblw = tblpr.find(qn("w:tblW"))
    if tblw is None:
        tblw = OxmlElement("w:tblW")
        tblpr.append(tblw)
    tblw.set(qn("w:w"), str(total))
    tblw.set(qn("w:type"), "dxa")

    tblind = tblpr.find(qn("w:tblInd"))
    if tblind is None:
        tblind = OxmlElement("w:tblInd")
        tblpr.append(tblind)
    tblind.set(qn("w:w"), str(indent_dxa))
    tblind.set(qn("w:type"), "dxa")

    layout = tblpr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblpr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        trpr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        trpr.append(cant_split)
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tcpr = cell._tc.get_or_add_tcPr()
            tcw = tcpr.find(qn("w:tcW"))
            if tcw is None:
                tcw = OxmlElement("w:tcW")
                tcpr.append(tcw)
            tcw.set(qn("w:w"), str(width))
            tcw.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)


def repeat_table_header(row):
    trpr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    trpr.append(tbl_header)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, text, fld_end])
    set_run_font(run, size=8.5, color=MUTED)


def add_hyperlink(paragraph, text_value: str, url: str, *, color=PURPLE, underline=True):
    part = paragraph.part
    rid = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rid)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), FONT)
    rpr.append(rfonts)
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    rpr.append(c)
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rpr.append(u)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "18")
    rpr.append(sz)
    run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text_value
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def create_numbering(doc: Document):
    numbering = doc.part.numbering_part.element
    existing_abs = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    existing_num = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    next_abs = max(existing_abs or [0]) + 1
    next_num = max(existing_num or [0]) + 1

    def add_definition(ordered: bool):
        nonlocal next_abs, next_num
        abstract_id = next_abs
        num_id = next_num
        next_abs += 1
        next_num += 1

        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        numfmt = OxmlElement("w:numFmt")
        numfmt.set(qn("w:val"), "decimal" if ordered else "bullet")
        lvl.append(numfmt)
        lvltext = OxmlElement("w:lvlText")
        lvltext.set(qn("w:val"), "%1." if ordered else "•")
        lvl.append(lvltext)
        jc = OxmlElement("w:lvlJc")
        jc.set(qn("w:val"), "left")
        lvl.append(jc)
        ppr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "540")
        tabs.append(tab)
        ppr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "540")
        ind.set(qn("w:hanging"), "280")
        ppr.append(ind)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:before"), "0")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "290")
        spacing.set(qn("w:lineRule"), "auto")
        ppr.append(spacing)
        lvl.append(ppr)
        rpr = OxmlElement("w:rPr")
        rfonts = OxmlElement("w:rFonts")
        for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
            rfonts.set(qn(f"w:{attr}"), FONT)
        rpr.append(rfonts)
        color = OxmlElement("w:color")
        color.set(qn("w:val"), PURPLE if not ordered else NAVY)
        rpr.append(color)
        lvl.append(rpr)
        abstract.append(lvl)
        numbering.append(abstract)

        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_ref = OxmlElement("w:abstractNumId")
        abstract_ref.set(qn("w:val"), str(abstract_id))
        num.append(abstract_ref)
        numbering.append(num)
        return num_id

    return add_definition(False), add_definition(True)


def apply_num(p, num_id: int):
    ppr = p._p.get_or_add_pPr()
    numpr = ppr.find(qn("w:numPr"))
    if numpr is None:
        numpr = OxmlElement("w:numPr")
        ppr.append(numpr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    nid = OxmlElement("w:numId")
    nid.set(qn("w:val"), str(num_id))
    numpr.extend([ilvl, nid])


def add_bullet(doc_or_cell, text_value: str, bullet_id: int, *, size=10.5, color=BODY, bold_prefix=None):
    p = doc_or_cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    apply_num(p, bullet_id)
    set_para_spacing(p, after=4, line=1.208, keep=True)
    if bold_prefix and text_value.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, size=size, color=NAVY, bold=True)
        r = p.add_run(text_value[len(bold_prefix):])
        set_run_font(r, size=size, color=color)
    else:
        r = p.add_run(text_value)
        set_run_font(r, size=size, color=color)
    return p


def add_numbered(doc_or_cell, text_value: str, number_id: int, *, size=10.5, color=BODY):
    p = doc_or_cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    apply_num(p, number_id)
    set_para_spacing(p, after=5, line=1.208, keep=True)
    r = p.add_run(text_value)
    set_run_font(r, size=size, color=color)
    return p


def clear_cell(cell):
    p = cell.paragraphs[0]
    for run in p.runs:
        run._element.getparent().remove(run._element)
    return p


def cell_text(cell, text_value: str, *, size=9.2, color=BODY, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, after=0):
    p = clear_cell(cell)
    p.alignment = align
    set_para_spacing(p, after=after, line=1.15, keep=True)
    r = p.add_run(text_value)
    set_run_font(r, size=size, color=color, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell, top=105, bottom=105, start=125, end=125)
    return p


def add_kicker(doc, text_value: str, *, align=WD_ALIGN_PARAGRAPH.LEFT, after=5):
    p = doc.add_paragraph(style="Kicker")
    p.alignment = align
    set_para_spacing(p, before=0, after=after, line=1.0, keep=True, keep_next=True)
    r = p.add_run(text_value.upper())
    set_run_font(r, size=8.3, color=PURPLE_SOFT, bold=True)
    return p


def add_heading(doc, text_value: str, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text_value)
    return p


def add_body(doc_or_cell, text_value: str, *, size=11, color=BODY, bold_prefix=None, italic=False, after=8, align=None, keep=False):
    p = doc_or_cell.add_paragraph()
    p.style = "Normal"
    if align is not None:
        p.alignment = align
    set_para_spacing(p, after=after, line=1.333, keep=keep)
    if bold_prefix and text_value.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, size=size, color=NAVY, bold=True)
        r = p.add_run(text_value[len(bold_prefix):])
        set_run_font(r, size=size, color=color, italic=italic)
    else:
        r = p.add_run(text_value)
        set_run_font(r, size=size, color=color, italic=italic)
    return p


def add_callout(doc, label: str, headline: str, body: str, *, fill=SURFACE, accent=PURPLE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_DXA], indent_dxa=220)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_borders(cell, color=accent, size=10, sides=("left",))
    set_cell_margins(cell, top=160, bottom=160, start=220, end=220)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = clear_cell(cell)
    set_para_spacing(p, after=3, line=1.0, keep=True, keep_next=True)
    r = p.add_run(label.upper())
    set_run_font(r, size=8, color=accent, bold=True)
    p = cell.add_paragraph()
    set_para_spacing(p, after=4, line=1.05, keep=True, keep_next=True)
    r = p.add_run(headline)
    set_run_font(r, size=16, color=NAVY, bold=True)
    p = cell.add_paragraph()
    set_para_spacing(p, after=0, line=1.22, keep=True)
    r = p.add_run(body)
    set_run_font(r, size=9.8, color=BODY)
    spacer = doc.add_paragraph()
    set_para_spacing(spacer, after=5)
    return table


def add_section_label(doc, number: str, text_value: str):
    p = doc.add_paragraph()
    set_para_spacing(p, before=0, after=8, line=1.0, keep=True, keep_next=True)
    r = p.add_run(f"{number}  ")
    set_run_font(r, size=8.5, color=PINK, bold=True)
    r = p.add_run(text_value.upper())
    set_run_font(r, size=8.5, color=PURPLE_SOFT, bold=True)
    return p


def add_page_break(doc):
    p = doc.add_paragraph()
    p.paragraph_format.page_break_before = True
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(1)
    r = p.add_run("\u200b")
    set_run_font(r, size=1, color=WHITE)


def add_table_spacer(doc, points=6):
    p = doc.add_paragraph()
    set_para_spacing(p, before=0, after=points)
    return p


def configure_styles(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(BODY)
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal._element.rPr.rFonts.set(qn("w:cs"), FONT)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    title = styles["Title"]
    title.font.name = FONT
    title.font.size = Pt(29)
    title.font.bold = True
    title.font.color.rgb = rgb(NAVY)
    title._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    title._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)
    title.paragraph_format.line_spacing = 1.02
    title.paragraph_format.keep_with_next = True

    subtitle = styles["Subtitle"]
    subtitle.font.name = FONT
    subtitle.font.size = Pt(13)
    subtitle.font.color.rgb = rgb(MUTED)
    subtitle._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    subtitle._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(14)
    subtitle.paragraph_format.line_spacing = 1.2

    heading_tokens = {
        1: (16, NAVY, 18, 10),
        2: (13, PURPLE_DARK, 12, 6),
        3: (12, NAVY_2, 8, 4),
    }
    for level, (size, color, before, after) in heading_tokens.items():
        st = styles[f"Heading {level}"]
        st.font.name = FONT
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = rgb(color)
        st._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        st._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.line_spacing = 1.1
        st.paragraph_format.keep_with_next = True
        st.paragraph_format.keep_together = True

    if "Kicker" not in styles:
        kicker = styles.add_style("Kicker", WD_STYLE_TYPE.PARAGRAPH)
    else:
        kicker = styles["Kicker"]
    kicker.font.name = FONT
    kicker.font.size = Pt(8.3)
    kicker.font.bold = True
    kicker.font.color.rgb = rgb(PURPLE_SOFT)
    kicker._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    kicker._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    kicker.paragraph_format.space_before = Pt(0)
    kicker.paragraph_format.space_after = Pt(5)
    kicker.paragraph_format.line_spacing = 1.0

    if "Small Text" not in styles:
        small = styles.add_style("Small Text", WD_STYLE_TYPE.PARAGRAPH)
    else:
        small = styles["Small Text"]
    small.font.name = FONT
    small.font.size = Pt(8.5)
    small.font.color.rgb = rgb(MUTED)
    small._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    small._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    small.paragraph_format.space_before = Pt(0)
    small.paragraph_format.space_after = Pt(4)
    small.paragraph_format.line_spacing = 1.15


def configure_section(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True


def add_header_footer(section):
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    set_para_spacing(p, after=0, line=1.0)
    r = p.add_run("INZINT OMAN  ·  CHATBOT.OM")
    set_run_font(r, size=8, color=NAVY, bold=True)
    r = p.add_run("\tPROPOSAL FOR TAIF AL-KAABI")
    set_run_font(r, size=8, color=MUTED, bold=True)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    set_para_spacing(p, after=0, line=1.0)
    r = p.add_run("Inzint LLC  ·  inzint.om  ·  contact@inzint.om")
    set_run_font(r, size=8, color=MUTED)
    r = p.add_run("\tCONFIDENTIAL  ·  ")
    set_run_font(r, size=8, color=MUTED, bold=True)
    add_page_number(p)

    first_footer = section.first_page_footer
    p = first_footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, after=0, line=1.0)
    r = p.add_run("Inzint LLC  ·  31, Building 13, 6125 Way, Muscat, Oman  ·  (+968) 7272 4832  ·  contact@inzint.om")
    set_run_font(r, size=8, color=MUTED)


def make_banner():
    TMP_DIR.mkdir(parents=True, exist_ok=True)
    canvas = Image.new("RGBA", (1800, 360), (0, 12, 96, 255))
    draw = ImageDraw.Draw(canvas, "RGBA")
    for i in range(720, 0, -8):
        alpha = max(0, int(65 * (1 - i / 720)))
        draw.ellipse((1320 - i, 180 - i, 1320 + i, 180 + i), fill=(98, 0, 238, alpha))
    draw.rectangle((0, 348, 1800, 360), fill=(255, 2, 102, 255))
    logo = Image.open(LOGO_PATH).convert("RGBA")
    ratio = 205 / logo.height
    logo = logo.resize((int(logo.width * ratio), 205), Image.Resampling.LANCZOS)
    canvas.alpha_composite(logo, (105, 74))
    canvas.convert("RGB").save(BANNER_PATH, quality=96)


def build_document():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    make_banner()
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    configure_section(section)
    add_header_footer(section)
    bullet_id, number_id = create_numbering(doc)

    props = doc.core_properties
    props.title = "WhatsApp AI Student Enquiry and Enrolment Support - SaaS Proposal"
    props.subject = "Indicative managed SaaS proposal for Taif Al-Kaabi and their training institute"
    props.author = "Inzint LLC"
    props.last_modified_by = "Inzint LLC"
    props.keywords = "Inzint Oman, Chatbot.OM, training institute, AI assistant, SaaS proposal"
    props.comments = "Prepared from verified product and implementation scope on 25 August 2026."
    props.created = datetime(2026, 8, 25, tzinfo=timezone.utc)
    props.modified = datetime(2026, 8, 25, tzinfo=timezone.utc)

    # Cover page -----------------------------------------------------------------
    banner = doc.add_picture(str(BANNER_PATH), width=Inches(6.5), height=Inches(1.3))
    banner._inline.docPr.set("title", "Inzint Oman")
    banner._inline.docPr.set(
        "descr",
        "Inzint Oman wordmark in white on the navy and purple brand banner.",
    )
    p = doc.paragraphs[-1]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, after=12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_para_spacing(p, before=0, after=18, line=1.25, keep=True)
    set_paragraph_shading(p, SURFACE)
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.18)
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(18)
    ppr = p._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    ppr.append(bidi)
    r = p.add_run("إذا رغبتم في استلام هذا العرض باللغة العربية، يرجى طلب ذلك، وسيسعدنا تزويدكم بنسخة عربية.")
    set_run_font(r, name=ARABIC_FONT, size=10.5, color=NAVY, bold=False)
    rpr = r._element.get_or_add_rPr()
    rtl = OxmlElement("w:rtl")
    rpr.append(rtl)

    add_kicker(doc, "Managed SaaS Proposal · 25 August 2026", align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("WhatsApp AI Student Enquiry\nand Enrolment Support")
    set_para_spacing(p, after=9, line=1.0, keep=True, keep_next=True)
    p = doc.add_paragraph(style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("A managed Chatbot.OM service proposal for course enquiries, lead capture, guidance, booking requests and follow-up.")
    set_para_spacing(p, after=22, line=1.2, keep=True)

    table = doc.add_table(rows=4, cols=2)
    set_table_geometry(table, [2160, 7200], indent_dxa=125)
    metadata = [
        ("Prepared for", "Taif Al-Kaabi and their Training Institute"),
        ("Prepared by", "Inzint LLC · Inzint Oman"),
        ("Proposal reference", "INZ-OM-TK-20260825"),
        ("Validity", "Indicative offer valid through 24 September 2026"),
    ]
    for idx, (label, value) in enumerate(metadata):
        for cell in table.rows[idx].cells:
            set_cell_no_borders(cell)
            set_cell_margins(cell, top=70, bottom=70, start=100, end=100)
        cell_text(table.cell(idx, 0), label.upper(), size=8.2, color=PURPLE_SOFT, bold=True)
        cell_text(table.cell(idx, 1), value, size=10, color=NAVY, bold=(idx == 0))
    add_table_spacer(doc, 11)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, after=0, line=1.15)
    r = p.add_run("AI and software, built for Oman.")
    set_run_font(r, size=10, color=PURPLE, bold=True)

    # Page 2 ---------------------------------------------------------------------
    add_page_break(doc)
    add_section_label(doc, "01", "Executive proposal")
    add_heading(doc, "A practical first step, with clear boundaries", 1)
    add_body(
        doc,
        "Dear Taif, thank you for sharing what your institute needs. The requested journey is clear: answer course questions quickly, guide each prospective student, capture their details, help them take a next step, and give your team a reliable record for follow-up.",
    )
    add_callout(
        doc,
        "Recommended operating plan",
        "Business · OMR 40 per month",
        "Includes 25,000 chat sessions per month; additional sessions are OMR 0.025 each. It is the recommended plan for an active institute intake cycle. Final activation remains subject to the WhatsApp readiness gate described below.",
        fill=SURFACE,
        accent=PURPLE,
    )
    add_heading(doc, "The proposal in four answers", 2)
    answers = [
        ("Price", "OMR 40/month recommended. Other paid options are shown on page 6. There is no separate discounted annual tariff in this offer."),
        ("Timing", "Configuration can begin immediately. WhatsApp activation normally depends on number and template approval; the final schedule will be confirmed after technical readiness testing."),
        ("Customization", "The monthly SaaS can be configured for branding, assistant name, institute data, questions, rules and supported workflows, but it is not a fully custom or customer-owned codebase."),
        ("Support", "Yes. Managed maintenance and remote support are included during an active subscription; response targets and service hours are fixed in the final order form."),
    ]
    table = doc.add_table(rows=len(answers), cols=2)
    set_table_geometry(table, [1800, 7560], indent_dxa=125)
    for idx, (label, body) in enumerate(answers):
        c0, c1 = table.rows[idx].cells
        set_cell_shading(c0, NAVY if idx == 0 else SURFACE_2)
        set_cell_shading(c1, SURFACE_2)
        set_cell_borders(c0, color=BORDER, size=5)
        set_cell_borders(c1, color=BORDER, size=5)
        cell_text(c0, label, size=9.2, color=WHITE if idx == 0 else NAVY, bold=True)
        cell_text(c1, body, size=9.1, color=BODY)
    add_table_spacer(doc, 6)
    add_callout(
        doc,
        "Deployment gate",
        "WhatsApp is confirmed before it is promised as live",
        "This indicative proposal does not assume an already-operational WhatsApp transport. Inzint will validate the proposed business number, provider path, inbound and outbound flows, and end-to-end message delivery before the final order form and go-live date. Any non-standard engineering will be separately identified and approved.",
        fill=AMBER_BG,
        accent=AMBER,
    )

    # Page 3 ---------------------------------------------------------------------
    add_page_break(doc)
    add_section_label(doc, "02", "Proposed student journey")
    add_heading(doc, "From first question to a staff-owned enrolment lead", 1)
    add_body(
        doc,
        "The assistant is designed to make the early part of enrolment immediate and consistent. It handles approved information and routine next steps; your staff keep control of final enrolment, payment and any decision that depends on judgement or live capacity.",
    )

    stages = [
        ("01", "ASK", "Courses, prices, dates, locations and conditions"),
        ("02", "GUIDE", "Match by goal, level, timing and preference"),
        ("03", "CAPTURE", "Name, phone, interest, intake and notes"),
        ("04", "BOOK", "Trial or consultation; record a seat request"),
        ("05", "HAND OVER", "Qualified lead, context and next action"),
    ]
    table = doc.add_table(rows=1, cols=5)
    set_table_geometry(table, [1872] * 5, indent_dxa=100)
    for idx, (num, title, body) in enumerate(stages):
        cell = table.cell(0, idx)
        set_cell_shading(cell, NAVY if idx == 0 else SURFACE)
        set_cell_borders(cell, color=WHITE if idx == 0 else BORDER, size=5)
        set_cell_margins(cell, top=130, bottom=130, start=100, end=100)
        p = clear_cell(cell)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_spacing(p, after=3, line=1.0, keep=True, keep_next=True)
        r = p.add_run(num)
        set_run_font(r, size=8, color=PINK if idx == 0 else PURPLE, bold=True)
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_spacing(p, after=5, line=1.0, keep=True, keep_next=True)
        r = p.add_run(title)
        set_run_font(r, size=8.7, color=WHITE if idx == 0 else NAVY, bold=True)
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_spacing(p, after=0, line=1.12, keep=True)
        r = p.add_run(body)
        set_run_font(r, size=8, color=WHITE if idx == 0 else BODY)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    add_table_spacer(doc, 8)

    add_heading(doc, "What the assistant can be taught", 2)
    for item in [
        "Course catalogue: descriptions, target audience, prerequisites and learning outcomes.",
        "Commercial information: approved prices, discounts, payment conditions and refund rules.",
        "Operations: intake dates, schedules, locations, class formats and contact details.",
        "Conversation rules: what it may answer, what it must not promise, and when it must hand over to a person.",
    ]:
        add_bullet(doc, item, bullet_id, size=10.3)

    add_heading(doc, "What your team can monitor", 2)
    add_body(
        doc,
        "The dashboard provides conversation logs, contact records, qualified leads, lead status, notes, assignees, knowledge sources, appointments, sentiment and activity trends. It is a conversation-and-lead dashboard, not a replacement for a full student information or institute-management system.",
        after=7,
    )
    add_callout(
        doc,
        "Safe guidance",
        "Conversational course recommendation, not an eligibility engine",
        "The assistant can recommend suitable courses from the approved catalogue and the student's stated goals. Formal eligibility, certification, credit transfer or complex prerequisite decisions remain with institute staff unless a dedicated rules engine is separately built.",
        fill=LILAC_BG,
        accent=PURPLE,
    )

    # Page 4 ---------------------------------------------------------------------
    add_page_break(doc)
    add_section_label(doc, "03", "Response to your requirements")
    add_heading(doc, "What is included, conditional or separately scoped", 1)
    add_body(
        doc,
        "The table below is the controlling scope summary for this proposal. A conditional or separate item is not included in the base monthly fee unless it is added to the final order form.",
        after=8,
    )

    rows = [
        ("WhatsApp Business + AI replies", "READINESS-GATED", "Proposed after the number, provider path, inbound replies and approved outbound templates pass end-to-end testing.", "Not represented as live before UAT; non-standard engineering may require a separate quote."),
        ("Institute knowledge", "INCLUDED", "Upload approved course, price, schedule, location and policy content from documents or a website.", "Institute supplies and approves the source material; the assistant is instructed not to invent missing facts."),
        ("Student data capture", "INCLUDED", "Capture configured fields and preserve the conversation; qualified enquiries become lead records with score, status, notes and owner.", "Anonymous or low-information conversations may remain logs rather than automatically becoming leads."),
        ("Course selection help", "INCLUDED", "Conversational guidance from the approved catalogue using goals, level, timing and preferences.", "No deterministic eligibility or prerequisite engine is included."),
        ("Booking and confirmations", "PARTIAL", "Book a trial or consultation in available appointment slots; standard confirmation is by email with a calendar invite.", "No native live course-seat capacity, final enrolment or standard WhatsApp confirmation/reminder."),
        ("Follow-up for non-registrants", "PARTIAL", "Lead stages, notes, staff ownership and follow-up tracking are available in the dashboard.", "Automatic identification and WhatsApp messaging of non-registrants is separate automation work."),
        ("Future CRM / institute system", "INTEGRATION-READY", "APIs, webhooks and Zapier events can carry leads, appointments and conversation events.", "Target-specific mapping, authentication, testing and bidirectional sync are separately scoped."),
        ("Management dashboard", "INCLUDED", "Monitor conversations, qualified leads, contacts, appointments, common questions, tasks and sentiment/activity trends.", "It is not a student information, attendance, payment or certificate-management system."),
    ]
    table = doc.add_table(rows=1, cols=4)
    set_table_geometry(table, [1980, 1620, 3240, 2520], indent_dxa=125)
    headers = ["REQUEST", "POSITION", "SAAS RESPONSE", "BOUNDARY"]
    for idx, text_value in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, NAVY)
        set_cell_borders(cell, color=WHITE, size=4)
        cell_text(cell, text_value, size=8, color=WHITE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    repeat_table_header(table.rows[0])
    status_style = {
        "INCLUDED": (GREEN_BG, GREEN),
        "READINESS-GATED": (AMBER_BG, AMBER),
        "PARTIAL": (AMBER_BG, AMBER),
        "INTEGRATION-READY": (LILAC_BG, PURPLE_DARK),
    }
    for ridx, (request, position, response, boundary) in enumerate(rows):
        cells = table.add_row().cells
        fill = WHITE if ridx % 2 == 0 else SURFACE_2
        for cell in cells:
            set_cell_shading(cell, fill)
            set_cell_borders(cell, color=BORDER, size=4)
        cell_text(cells[0], request, size=8.25, color=NAVY, bold=True)
        pfill, pcolor = status_style[position]
        set_cell_shading(cells[1], pfill)
        cell_text(cells[1], position, size=7.4, color=pcolor, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        cell_text(cells[2], response, size=8.05, color=BODY)
        cell_text(cells[3], boundary, size=7.9, color=MUTED)
    set_table_geometry(table, [1980, 1620, 3240, 2520], indent_dxa=125)
    add_table_spacer(doc, 5)
    p = doc.add_paragraph(style="Small Text")
    set_para_spacing(p, before=4, after=0, line=1.15)
    r = p.add_run("Scope note: ")
    set_run_font(r, size=8.5, color=NAVY, bold=True)
    r = p.add_run("Where the institute later adds a compatible CRM or management system, the current platform can be used as the conversation layer and connected through a separately approved integration scope.")
    set_run_font(r, size=8.5, color=MUTED)

    # Page 5 ---------------------------------------------------------------------
    add_page_break(doc)
    add_section_label(doc, "04", "Customization and ownership")
    add_heading(doc, "Monthly SaaS is configurable, not fully custom", 1)
    add_body(
        doc,
        "Under the monthly model, Inzint operates the software and gives the institute access to a managed service. Configuration is intentionally bounded so upgrades and maintenance remain manageable across subscribers.",
    )
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [4680, 4680], indent_dxa=190)
    left, right = table.rows[0].cells
    for cell in (left, right):
        set_cell_margins(cell, top=180, bottom=180, start=190, end=190)
        set_cell_borders(cell, color=BORDER, size=6)
    set_cell_shading(left, SURFACE)
    set_cell_shading(right, NAVY)

    p = clear_cell(left)
    set_para_spacing(p, after=4, line=1.0, keep=True, keep_next=True)
    r = p.add_run("MONTHLY SAAS")
    set_run_font(r, size=8, color=PURPLE, bold=True)
    p = left.add_paragraph()
    set_para_spacing(p, after=9, line=1.05, keep=True, keep_next=True)
    r = p.add_run("Configuration included")
    set_run_font(r, size=15, color=NAVY, bold=True)
    for item in [
        "Assistant name, greeting, tone and system instructions",
        "Brand colours, theme, launcher text/icon and welcome message",
        "Course, fee, timetable, location, condition and policy data",
        "Lead fields, qualification questions and conversation guardrails",
        "Working hours, timezone, appointment duration and basic team access",
    ]:
        add_bullet(left, item, bullet_id, size=9.1, color=BODY)

    p = clear_cell(right)
    set_para_spacing(p, after=4, line=1.0, keep=True, keep_next=True)
    r = p.add_run("SEPARATE PROPOSAL")
    set_run_font(r, size=8, color=PINK, bold=True)
    p = right.add_paragraph()
    set_para_spacing(p, after=9, line=1.05, keep=True, keep_next=True)
    r = p.add_run("Owned copy / custom software")
    set_run_font(r, size=15, color=WHITE, bold=True)
    for item in [
        "Custom student, course, cohort, seat-capacity and enrolment modules",
        "Deep CRM or institute-management workflows and bespoke integrations",
        "Custom data model, roles, dashboards and approval processes",
        "Customer-controlled deployment, release cadence and security architecture",
        "Source-code handover, ownership/licence rights and long-term roadmap control",
    ]:
        p = right.add_paragraph()
        apply_num(p, bullet_id)
        set_para_spacing(p, after=4, line=1.208, keep=True)
        r = p.add_run(item)
        set_run_font(r, size=9.1, color=WHITE)
    add_table_spacer(doc, 8)

    add_callout(
        doc,
        "Direct answer to question 3",
        "Limited customization is included; full control requires an owned solution",
        "If Taif Al-Kaabi would like unrestricted customization and control of the software, Inzint will be happy to prepare a separate one-time custom-software proposal. That option is not priced or included in this monthly SaaS offer.",
        fill=LILAC_BG,
        accent=PURPLE,
    )
    add_heading(doc, "Changes during the subscription", 2)
    add_body(
        doc,
        "Routine updates to approved content and reasonable in-scope configuration changes can be handled through support. New transports, modules, data structures, third-party connectors or material workflow redesigns are assessed and quoted separately before work begins.",
        after=0,
    )

    # Page 6 ---------------------------------------------------------------------
    add_page_break(doc)
    add_section_label(doc, "05", "Commercial offer")
    add_heading(doc, "Choose by conversation volume", 1)
    add_body(
        doc,
        "The paid plans below share the standard product capability set. The published differences are the monthly session allowance and overage rate. Technical activation and any separately scoped integration work remain governed by the final order form.",
    )
    add_callout(
        doc,
        "Recommended",
        "Business · OMR 40/month",
        "25,000 sessions each month, then OMR 0.025 per additional session. This gives the institute enough headroom for an active intake period without moving directly to an unlimited plan.",
        fill=SURFACE,
        accent=PURPLE,
    )

    pricing = [
        ("Starter", "OMR 10", "2,000", "OMR 0.050", "Limited preview only; not recommended for a live enrolment cycle."),
        ("Business", "OMR 40", "25,000", "OMR 0.025", "Recommended for normal institute operations."),
        ("Enterprise", "OMR 135", "Unlimited", "None", "High-volume use where session overage must be eliminated."),
    ]
    table = doc.add_table(rows=1, cols=5)
    set_table_geometry(table, [1440, 1440, 1440, 1440, 3600], indent_dxa=125)
    headers = ["PLAN", "MONTHLY", "SESSIONS", "EXTRA SESSION", "POSITIONING"]
    for idx, text_value in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, NAVY)
        set_cell_borders(cell, color=WHITE, size=4)
        cell_text(cell, text_value, size=7.8, color=WHITE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    repeat_table_header(table.rows[0])
    for ridx, row in enumerate(pricing):
        cells = table.add_row().cells
        highlight = ridx == 1
        for cell in cells:
            set_cell_shading(cell, LILAC_BG if highlight else (WHITE if ridx % 2 == 0 else SURFACE_2))
            set_cell_borders(cell, color=PURPLE if highlight else BORDER, size=6 if highlight else 4)
        for cidx, value in enumerate(row):
            cell_text(
                cells[cidx],
                value,
                size=8.5 if cidx < 4 else 8.2,
                color=NAVY if highlight or cidx == 0 else BODY,
                bold=(highlight or cidx == 0),
                align=WD_ALIGN_PARAGRAPH.CENTER if cidx < 4 else WD_ALIGN_PARAGRAPH.LEFT,
            )
    set_table_geometry(table, [1440, 1440, 1440, 1440, 3600], indent_dxa=125)
    add_table_spacer(doc, 7)

    add_heading(doc, "Annual pricing", 2)
    add_body(
        doc,
        "No separate discounted annual tariff is included in this offer. For budgeting only, twelve months at the current base monthly rates equal OMR 120 for Starter, OMR 480 for Business and OMR 1,620 for Enterprise, before any additional-session usage or separately approved work.",
        after=7,
    )
    add_heading(doc, "Commercial notes", 2)
    for item in [
        "All amounts are stated in Omani Rials (OMR).",
        "Usage is measured by the platform session counter; the session boundary will be defined in the final order form.",
        "No implementation fee, tax treatment or third-party platform charge is assumed in the amounts above; any applicable item will be confirmed before order acceptance.",
        "The subscription provides a right to use the managed SaaS service. It does not transfer source code, infrastructure or ownership rights.",
    ]:
        add_bullet(doc, item, bullet_id, size=9.7)

    # Page 7 ---------------------------------------------------------------------
    add_page_break(doc)
    add_section_label(doc, "06", "Implementation and support")
    add_heading(doc, "A short configuration cycle, then a controlled launch", 1)
    add_body(
        doc,
        "The standard configuration work can move quickly once the institute's information and access are ready. WhatsApp is the schedule dependency because its number, provider path and outbound templates must be approved and tested.",
    )

    implementation = [
        ("1", "READINESS", "Confirm channel path, number status, consent basis, scope and selected plan."),
        ("2", "TEACH", "Load approved course, fee, schedule, location, policy and FAQ content."),
        ("3", "CONFIGURE", "Set assistant name, tone, questions, lead fields, handover rules and appearance."),
        ("4", "TEST", "Run representative Arabic/English questions, lead capture, booking and handover checks."),
        ("5", "LAUNCH", "Approve go-live, monitor early conversations and refine in-scope content."),
    ]
    table = doc.add_table(rows=len(implementation), cols=3)
    set_table_geometry(table, [720, 1800, 6840], indent_dxa=125)
    for idx, (num, title, body) in enumerate(implementation):
        c0, c1, c2 = table.rows[idx].cells
        for cell in (c0, c1, c2):
            set_cell_borders(cell, color=BORDER, size=4)
        set_cell_shading(c0, NAVY)
        set_cell_shading(c1, SURFACE)
        set_cell_shading(c2, WHITE if idx % 2 == 0 else SURFACE_2)
        cell_text(c0, num, size=10, color=WHITE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        cell_text(c1, title, size=8.1, color=PURPLE_DARK, bold=True)
        cell_text(c2, body, size=8.9, color=BODY)
    add_table_spacer(doc, 7)

    add_heading(doc, "Indicative duration", 2)
    add_callout(
        doc,
        "Timing answer",
        "Typically a few days after complete inputs",
        "Inzint starts the technical and template readiness work on day one. The committed go-live date will be issued after the proposed WhatsApp number and messaging flow pass readiness checks; third-party approval timing is outside Inzint's direct control.",
        fill=AMBER_BG,
        accent=AMBER,
    )

    add_heading(doc, "Support and maintenance after launch", 2)
    add_body(
        doc,
        "Yes. During an active subscription, Inzint provides managed platform maintenance, routine fixes and updates, incident support, and reasonable assistance with in-scope content or configuration changes. The final order form will state support channels, business hours, response targets and any enhanced-service option.",
        after=5,
    )
    add_body(
        doc,
        "Separately charged work includes new product features, major workflow redesign, custom integrations, data migration or cleanup, customer-specific infrastructure, content authoring, and owned-software development.",
        size=9.6,
        color=MUTED,
        after=0,
    )

    # Page 8 ---------------------------------------------------------------------
    add_page_break(doc)
    add_section_label(doc, "07", "Assumptions and next steps")
    add_heading(doc, "What we need from the institute", 1)
    for item in [
        "An approved course catalogue, current prices, schedules, locations, terms and frequent questions.",
        "A decision on which WhatsApp Business number is proposed and confirmation of its present account status.",
        "Approved consent language and proposed outbound template wording for confirmations or follow-up.",
        "A named institute owner for answers, testing, lead routing and final go-live approval.",
        "Details of any CRM, calendar or institute-management system that may be connected later.",
    ]:
        add_bullet(doc, item, bullet_id, size=10.1)

    add_heading(doc, "Important boundaries", 2)
    boundaries = [
        ("ENROLMENT", "Final enrolment, payment and live course-seat capacity remain with staff unless a separate system integration or custom module is approved."),
        ("MESSAGING", "Automated outbound WhatsApp confirmations, reminders and non-registrant follow-up are not included until the WhatsApp channel and automation are expressly activated in the order form."),
        ("INTEGRATIONS", "The platform is integration-ready, but target-system mapping, authentication, testing and support are not turnkey inclusions."),
        ("HOSTING", "This indicative SaaS offer does not promise Oman-only or customer-controlled hosting. Any residency, on-premises or private-cloud requirement needs a separate deployment and security scope."),
        ("CONTENT", "The institute is responsible for the accuracy, authority and timely update of course, price, schedule and policy information supplied to the assistant."),
    ]
    table = doc.add_table(rows=len(boundaries), cols=2)
    set_table_geometry(table, [1800, 7560], indent_dxa=125)
    for idx, (label, body) in enumerate(boundaries):
        c0, c1 = table.rows[idx].cells
        set_cell_shading(c0, NAVY if idx == 0 else SURFACE)
        set_cell_shading(c1, WHITE if idx % 2 == 0 else SURFACE_2)
        set_cell_borders(c0, color=BORDER, size=4)
        set_cell_borders(c1, color=BORDER, size=4)
        cell_text(c0, label, size=8, color=WHITE if idx == 0 else PURPLE_DARK, bold=True)
        cell_text(c1, body, size=8.5, color=BODY)
    add_table_spacer(doc, 7)

    add_heading(doc, "To proceed", 2)
    for item in [
        "Confirm the preferred plan. Business is the recommended operational option; Starter is intended only for a limited preview.",
        "Share the proposed WhatsApp number and the institute information listed above.",
        "Complete a short readiness and scope review with Inzint Oman at (+968) 7272 4832 or contact@inzint.om.",
        "Receive and approve the final order form, including activated channels, implementation items, support terms and go-live date.",
    ]:
        add_numbered(doc, item, number_id, size=9.8)

    # Document-level settings and compatibility ---------------------------------
    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")
    compat = settings.find(qn("w:compat"))
    if compat is None:
        compat = OxmlElement("w:compat")
        settings.append(compat)
    setting = OxmlElement("w:compatSetting")
    setting.set(qn("w:name"), "compatibilityMode")
    setting.set(qn("w:uri"), "http://schemas.microsoft.com/office/word")
    setting.set(qn("w:val"), "15")
    compat.append(setting)

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build_document()
